import os
import time
import re
import logging
from copy import deepcopy
from groq import Groq
from docx import Document
from deep_translator import GoogleTranslator
import uuid
import hashlib

os.environ["GROQ_API_KEY"] = "get_ur_own_key_bitch"

# Configure logging to print to CLI
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# Create a simple translation cache to avoid re-translating identical content
translation_cache = {}

def translate_text_google(text, target_language='es'):
    """Translate text using Google Translate while handling API limits."""
    if not text.strip():
        return text
    if re.match(r'^[\d\W\s]+$', text):
        return text
    
    # Check cache first
    cache_key = f"google_{target_language}_{hashlib.md5(text.encode()).hexdigest()}"
    if cache_key in translation_cache:
        logging.info("Google: Using cached translation for text")
        return translation_cache[cache_key]
        
    try:
        logging.info("Google: Translating text chunk with length %d", len(text))
        time.sleep(0.2)  # Prevent API rate limit errors
        if len(text) > 5000:
            chunks = []
            for i in range(0, len(text), 4000):
                chunk = text[i:i+4000]
                translated_chunk = GoogleTranslator(source='auto', target=target_language).translate(chunk)
                chunks.append(translated_chunk)
            result = ''.join(chunks)
        else:
            result = GoogleTranslator(source='auto', target=target_language).translate(text)
        
        # Cache the result
        translation_cache[cache_key] = result
        return result
    except Exception as e:
        logging.error("Google Translation error: %s", e)
        return text

def translation_using_groq(texts, target_language='hi'):
    """
    Efficiently translate texts using Groq API with intelligent batching.
    This function takes a list of text chunks and translates them in optimal batches.
    """
    # Filter out empty texts or those with only special characters
    texts_to_translate = []
    original_indices = []
    
    for i, text in enumerate(texts):
        if not text.strip() or re.match(r'^[\d\W\s]+$', text):
            texts[i] = text  # Keep unchanged
        else:
            # Check cache first
            cache_key = f"groq_{target_language}_{hashlib.md5(text.encode()).hexdigest()}"
            if cache_key in translation_cache:
                logging.info(f"Groq: Using cached translation for text chunk {i+1}")
                texts[i] = translation_cache[cache_key]
            else:
                texts_to_translate.append(text)
                original_indices.append(i)
    
    if not texts_to_translate:
        return texts  # All texts were either empty, special chars, or found in cache
    
    try:
        client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
        
        # Create a system prompt for high-quality translations
        system_prompt = (
            f"You are a professional translator specialized in {target_language}. "
            f"Your task is to translate the following text from English to {target_language}. "
            "Translate natural-sounding text that maintains the original meaning and tone. "
            "Provide ONLY the translated text without any additional comments, explanations, or formatting."
        )
        
        # Define optimal batch sizes (30K is a good size for Mistral without hitting context limits)
        optimal_batch_size = 5000
        
        # Group small texts together to minimize API calls
        batches = []
        current_batch = []
        current_batch_size = 0
        current_batch_indices = []
        
        for i, text in enumerate(texts_to_translate):
            # If adding this text exceeds the optimal batch size, finalize the current batch
            if current_batch_size + len(text) > optimal_batch_size and current_batch:
                batches.append({
                    'texts': current_batch, 
                    'indices': current_batch_indices
                })
                current_batch = []
                current_batch_size = 0
                current_batch_indices = []
            
            current_batch.append(text)
            current_batch_size += len(text)
            current_batch_indices.append(original_indices[i])
        
        # Add the last batch if it's not empty
        if current_batch:
            batches.append({
                'texts': current_batch, 
                'indices': current_batch_indices
            })
        
        logging.info(f"Groq: Optimized {len(texts_to_translate)} text chunks into {len(batches)} API calls")
        
        # Process each batch
        for batch_idx, batch in enumerate(batches):
            # Generate unique markers to identify each text in the batch
            markers = []
            batch_text = ""
            
            for text in batch['texts']:
                marker = f"<<<ITEM_{uuid.uuid4()}>>>"
                markers.append(marker)
                batch_text += f"{marker}\n{text}\n\n"
            
            # Create the prompt with instructions to keep the markers
            user_prompt = (
                f"Translate the following texts from English to {target_language}. "
                "Each text is marked with a unique marker. Keep these markers in your response "
                "exactly as they are, and provide the translation immediately after each marker.\n\n"
                f"{batch_text}"
            )
            
            # Translate the batch with retry logic
            max_retries = 3
            success = False
            
            for attempt in range(max_retries):
                try:
                    logging.info(f"Groq: Translating batch {batch_idx+1}/{len(batches)} (attempt {attempt+1})")
                    
                    messages = [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ]
                    
                    chat_completion = client.chat.completions.create(
                        messages=messages,
                        model="mistral-saba-24b",
                        temperature=0.1,  # Low temperature for consistent translations
                    )
                    
                    response = chat_completion.choices[0].message.content.strip()
                    success = True
                    break
                except Exception as e:
                    logging.error(f"Groq API error (attempt {attempt+1}/{max_retries}): {e}")
                    if attempt < max_retries - 1:
                        # Exponential backoff
                        sleep_time = 2 ** attempt
                        logging.info(f"Retrying in {sleep_time} seconds...")
                        time.sleep(sleep_time)
            
            # If all translation attempts failed, use original texts
            if not success:
                logging.warning(f"All translation attempts failed for batch {batch_idx+1}")
                for i, idx in enumerate(batch['indices']):
                    texts[idx] = batch['texts'][i]
                continue
            
            # Extract translated texts using the markers
            for i, marker in enumerate(markers):
                if marker in response:
                    # Extract the text between this marker and the next one (or end of response)
                    start_idx = response.find(marker) + len(marker)
                    if i < len(markers) - 1 and markers[i+1] in response:
                        end_idx = response.find(markers[i+1])
                    else:
                        end_idx = len(response)
                    
                    translated_text = response[start_idx:end_idx].strip()
                    
                    # Store in the cache
                    original_text = batch['texts'][i]
                    cache_key = f"groq_{target_language}_{hashlib.md5(original_text.encode()).hexdigest()}"
                    translation_cache[cache_key] = translated_text
                    
                    # Update the result array
                    texts[batch['indices'][i]] = translated_text
                else:
                    logging.warning(f"Marker {marker} not found in response; using original text")
                    texts[batch['indices'][i]] = batch['texts'][i]
            
            # Add a delay between batch API calls to avoid rate limits
            if batch_idx < len(batches) - 1:
                time.sleep(1)
        
        return texts
        
    except Exception as e:
        logging.error(f"Groq Translation error: {e}")
        # Return original texts for any that haven't been translated
        for i, idx in enumerate(original_indices):
            if idx < len(texts) and texts[idx] == texts_to_translate[i]:
                texts[idx] = texts_to_translate[i]
        return texts

def extract_hyperlinks(run):
    """Extract hyperlinks from a run and return a placeholder."""
    for child in run._element:
        if child.tag.endswith('hyperlink'):
            return {'element': deepcopy(child), 'text': run.text}
    return None

def run_has_image(run):
    """Detect if a run contains an image."""
    return any('drawing' in child.tag for child in run._element)

def preserve_formatting_with_placeholders(paragraph):
    """Extracts text, hyperlinks, images, and formatting."""
    text_parts, format_map, hyperlinks, images = [], [], [], []
    for run in paragraph.runs:
        if run_has_image(run):
            logging.info("Found an image in a run; preserving its element.")
            images.append({'position': len(text_parts), 'element': deepcopy(run._element)})
            format_map.append(None)
            continue
        
        hyperlink_data = extract_hyperlinks(run)
        if hyperlink_data:
            logging.info("Found a hyperlink in a run; preserving it.")
            hyperlinks.append({
                'position': len(text_parts),
                'element': hyperlink_data['element'],
                'text': hyperlink_data['text']
            })
            text_parts.append(f"{{HYPERLINK_{len(hyperlinks)}}}")
            format_map.append(None)
        else:
            text_parts.append(run.text)
            format_map.append({
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size,
                'font_name': run.font.name,
                'color': run.font.color.rgb if run.font.color else None,
                'highlight_color': run.font.highlight_color
            })
    return text_parts, format_map, hyperlinks, images

def apply_formatted_text(paragraph, text_parts, format_map, hyperlinks, images):
    """Applies translated text with formatting, restoring hyperlinks and images."""
    paragraph.clear()
    for text, fmt in zip(text_parts, format_map):
        run = paragraph.add_run(text)
        if fmt:
            run.bold = fmt['bold']
            run.italic = fmt['italic']
            run.underline = fmt['underline']
            if fmt['font_size']:
                run.font.size = fmt['font_size']
            if fmt['font_name']:
                run.font.name = fmt['font_name']
            if fmt['color']:
                run.font.color.rgb = fmt['color']
            if fmt['highlight_color']:
                run.font.highlight_color = fmt['highlight_color']
    
    for hyperlink in hyperlinks:
        paragraph._element.append(hyperlink['element'])
    
    for image in images:
        paragraph._element.append(image['element'])

def translate_paragraph_advanced(paragraph, target_language, engine='google'):
    """Translates a paragraph while preserving formatting, hyperlinks, and images."""
    logging.info("Translating paragraph with %d runs", len(paragraph.runs))
    text_parts, format_map, hyperlinks, images = preserve_formatting_with_placeholders(paragraph)
    
    # Find hyperlink placeholders and actual hyperlink texts for translation
    hyperlink_texts = []
    for i, part in enumerate(text_parts):
        for hyperlink in hyperlinks:
            if part == f"{{HYPERLINK_{hyperlink['position'] + 1}}}":
                hyperlink_texts.append(hyperlink['text'])
    
    # Choose translation engine
    if engine.lower() == 'groq':
        # Translate all paragraph text parts in one optimized batch
        translated_parts = translation_using_groq(text_parts, target_language)
        
        # Also translate hyperlink texts if any
        if hyperlink_texts:
            translated_hyperlinks = translation_using_groq(hyperlink_texts, target_language)
            
            # Update hyperlink texts with translations
            for i, hyperlink in enumerate(hyperlinks):
                if i < len(translated_hyperlinks):
                    hyperlink['text'] = translated_hyperlinks[i]
    else:
        # Google translate each part individually
        translated_parts = [translate_text_google(part, target_language) for part in text_parts]
        
        # Also translate hyperlink texts
        for hyperlink in hyperlinks:
            hyperlink['text'] = translate_text_google(hyperlink['text'], target_language)
    
    apply_formatted_text(paragraph, translated_parts, format_map, hyperlinks, images)

def collect_all_paragraphs(doc):
    """Collect all paragraphs from document, including those in tables and sections."""
    all_paragraphs = []
    
    # Main document paragraphs
    all_paragraphs.extend(doc.paragraphs)
    
    # Table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    
    # Header and footer paragraphs if needed
    for section in doc.sections:
        for header in [section.header]:
            all_paragraphs.extend(header.paragraphs)
        for footer in [section.footer]:
            all_paragraphs.extend(footer.paragraphs)
    
    return all_paragraphs

def translate_table_advanced(table, target_language, engine='google'):
    """Translates text inside a table while preserving structure."""
    logging.info("Translating a table with %d rows", len(table.rows))
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                translate_paragraph_advanced(paragraph, target_language, engine)

def translate_docx_advanced(input_file, output_file, target_language='es', engine='google'):
    """Translates a DOCX file while preserving formatting, hyperlinks, images, and tables."""
    logging.info("Starting document translation: %s -> %s", input_file, output_file)
    doc = Document(input_file)
    
    # Collect all document paragraphs
    all_paragraphs = collect_all_paragraphs(doc)
    logging.info(f"Document has {len(all_paragraphs)} total paragraphs to translate")
    
    # Translate document paragraphs
    for idx, paragraph in enumerate(doc.paragraphs):
        logging.info("Translating document paragraph %d/%d", idx + 1, len(doc.paragraphs))
        translate_paragraph_advanced(paragraph, target_language, engine)
    
    # Translate tables
    for tidx, table in enumerate(doc.tables):
        logging.info("Translating table %d/%d", tidx + 1, len(doc.tables))
        translate_table_advanced(table, target_language, engine)
    
    # Translate headers and footers
    for section in doc.sections:
        for header in [section.header]:
            for paragraph in header.paragraphs:
                translate_paragraph_advanced(paragraph, target_language, engine)
        
        for footer in [section.footer]:
            for paragraph in footer.paragraphs:
                translate_paragraph_advanced(paragraph, target_language, engine)
    
    # Save the translated document
    doc.save(output_file)
    logging.info("Translation completed using %s engine. Saved to %s", engine.capitalize(), output_file)

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(
        description='Translate DOCX files while preserving formatting, hyperlinks, images, and tables.'
    )
    parser.add_argument('input_file', help='Input DOCX file path')
    parser.add_argument('--output_file', help='Output DOCX file path')
    parser.add_argument('--target_language', default='es', help='Target language code (default: es for Spanish)')
    parser.add_argument('--engine', default='google', choices=['google', 'groq'], help='Translation engine to use (default: google)')
    args = parser.parse_args()
    if not args.output_file:
        args.output_file = f"{os.path.splitext(args.input_file)[0]}_{args.target_language}.docx"
    translate_docx_advanced(args.input_file, args.output_file, args.target_language, args.engine)
